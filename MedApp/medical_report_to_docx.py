import os
import base64
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class MedicalReport():

    storage_path = os.getcwd() + '/temp_storage/'

    @staticmethod
    def docx_to_base64(name):
        data = open(name, "rb").read()
        b64 = base64.b64encode(data)
        return b64.decode('utf-8')

    @staticmethod
    def create_report(doctor, patient, diagnosis, photo):
        name = str(patient.pk) + '_' + patient.last_name + ".docx"
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        par1 = doc.add_paragraph()
        par1.add_run('Медицинское заключение.').bold = True
        par1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        par2 = doc.add_paragraph()
        par2.add_run('Врач, проводивший исследование: \n').bold = True
        par2.add_run(doctor.last_name + ' ' + doctor.first_name + ' ' + doctor.middle_name)

        par3 = doc.add_paragraph()
        par3.add_run('Пациент: ').bold = True
        par3.add_run(patient.last_name + ' ' + patient.first_name + ' ' + patient.middle_name + '\n')
        par3.add_run('Год рождения: ').bold = True
        par3.add_run(str(patient.date_of_birth))

        par4 = doc.add_paragraph()
        par4.add_run('Диагноз: ').bold = True
        par4.add_run(diagnosis + '\n')
        par4.add_run('Дата проведения исследования: ').bold = True
        par4.add_run(str(photo.date_of_research).split('.')[0] + '\n')
        par4.add_run('Дата создания снимка: ').bold = True
        par4.add_run(str(photo.date_of_creation).split('+')[0])

        par5 = doc.add_paragraph()
        par5.add_run('Снимок').bold = True
        par5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_picture(MedicalReport.storage_path + photo.photo + '.jpeg', width = Cm(13))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(name)
        return name
